import os
import json
from typing import List, Dict, Any, Optional
import datetime
import logging
from pathlib import Path
import markdown2
from docx import Document

from crewai import Agent, Task, Crew, Process
from langchain_openai import ChatOpenAI
from crewai.tasks.task_output import TaskOutput
from pydantic import Field

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[logging.StreamHandler()]
)
logger = logging.getLogger("startup_research")

# Ensure environment variables are loaded
if not os.getenv("OPENAI_API_KEY"):
    logger.warning("OPENAI_API_KEY not found in environment variables")

class StartupResearchCrew:
    """
    A crew of AI agents that work together to research and evaluate startup ideas.
    """
    
    def __init__(
        self,
        model_name: str = "gpt-3.5-turbo",
        temperature: float = 0.5,
        process: Process = Process.sequential
    ):
        """
        Initialize the research crew.
        
        Args:
            model_name: The name of the LLM model to use
            temperature: The temperature setting for the LLM
            process: The process type for the crew
        """
        self.model_name = model_name
        self.temperature = temperature
        self.process = process
        
        # Initialize the LLM
        self.llm = ChatOpenAI(
            model_name=model_name,
            temperature=temperature
        )
        
        # Create output directory
        self.output_dir = Path("research_outputs")
        self.output_dir.mkdir(exist_ok=True)
        
        logger.info(f"StartupResearchCrew initialized with {model_name} at temp {temperature}")

    def create_agents(self) -> Dict[str, Agent]:
        """Create and return all the agents for the research crew"""
        
        # Planner agent - coordinates the research plan
        planner = Agent(
            role="Startup Research Planner",
            goal="Create a comprehensive research plan for evaluating a startup idea",
            backstory="""You are an experienced startup strategist who has helped 
            numerous founders evaluate their ideas. You excel at breaking down complex 
            research questions into clear, actionable tasks.""",
            verbose=True,
            allow_delegation=True,
            llm=self.llm
        )
        
        # Researcher agent - gathers market data and trends
        researcher = Agent(
            role="Market Research Specialist",
            goal="Gather comprehensive, factual information about the market, trends, and competitors",
            backstory="""You are a diligent market researcher with expertise in analyzing 
            emerging industries. You have a knack for finding relevant data and identifying 
            key market trends that others might miss.""",
            verbose=True,
            llm=self.llm
        )
        
        # Fact checker agent - validates research findings
        fact_checker = Agent(
            role="Research Fact Checker",
            goal="Verify information accuracy and ensure research is based on reliable sources",
            backstory="""You are a meticulous fact-checker with a background in journalism 
            and academic research. You have a critical eye for distinguishing between 
            reliable information and speculation.""",
            verbose=True,
            llm=self.llm
        )
        
        # Summarizer agent - condenses findings into clear points
        summarizer = Agent(
            role="Research Summarizer",
            goal="Condense complex information into clear, concise bullet points and key insights",
            backstory="""You excel at distilling complex information into easily 
            digestible summaries. You can identify the most important points in any research 
            and present them in a clear, structured way.""",
            verbose=True,
            llm=self.llm
        )
        
        # Analyst agent - performs SWOT and critical evaluation
        analyst = Agent(
            role="Startup Business Analyst",
            goal="Critically evaluate the business potential of the startup idea",
            backstory="""You have analyzed hundreds of startups across various industries. 
            You specialize in SWOT analysis, identifying competitive advantages, and 
            assessing business model viability.""",
            verbose=True,
            llm=self.llm
        )
        
        # Strategy agent - recommends go-to-market approach
        strategist = Agent(
            role="Go-to-Market Strategist",
            goal="Develop effective product positioning and go-to-market recommendations",
            backstory="""You're a seasoned go-to-market expert who has helped numerous 
            startups successfully launch their products. You know how to identify the 
            right channels, positioning, and business models for new ventures.""",
            verbose=True,
            llm=self.llm
        )
        
        # Writer agent - creates the final report
        writer = Agent(
            role="Startup Report Writer",
            goal="Create a clear, comprehensive report that communicates all findings",
            backstory="""You are an expert business writer who specializes in creating 
            engaging, insightful reports. You know how to structure information for 
            maximum clarity and impact, with executive-friendly language.""",
            verbose=True,
            llm=self.llm
        )
        
        # Critic agent - reviews and improves the report
        critic = Agent(
            role="Report Critic and Editor",
            goal="Improve the clarity, logic and presentation of the research report",
            backstory="""You have edited hundreds of business reports and presentations.
            You have a keen eye for logical inconsistencies, clarity issues, and areas
            where additional evidence or explanation would strengthen the argument.""",
            verbose=True,
            llm=self.llm
        )
        
        return {
            "planner": planner,
            "researcher": researcher,
            "fact_checker": fact_checker,
            "summarizer": summarizer, 
            "analyst": analyst,
            "strategist": strategist,
            "writer": writer,
            "critic": critic
        }
        
    def create_tasks(self, agents: Dict[str, Agent], startup_idea: str) -> List[Task]:
        """Create the sequence of research tasks"""
        
        planning_task = Task(
            description=f"""
            Create a comprehensive plan to evaluate the startup idea: '{startup_idea}'.
            Break down the research process into key areas to investigate, including:
            1. Market size and growth potential
            2. Target customer segments and pain points
            3. Competitive landscape and differentiation
            4. Business model and monetization options
            5. Key risks and challenges
            
            For each area, specify what exact information we need to gather and why.
            The plan should be detailed enough to guide our research but focused on
            the most relevant aspects for evaluating this specific startup idea.
            """,
            agent=agents["planner"],
            expected_output="A structured research plan with 5-7 key areas to investigate",
            output_file="research_plan.md"
        )
        
        research_task = Task(
            description=f"""
            Research the startup idea: '{startup_idea}' based on the research plan.
            Focus on gathering factual information about:
            - Market size, growth rates, and trends
            - Key players and competitors in this space
            - Target customer segments and their needs
            - Similar solutions and their business models
            - Relevant technologies and their maturity
            - Regulatory considerations or compliance requirements
            
            Use search tools to gather information from credible sources.
            Include specific data points, statistics, and example companies whenever possible.
            Cite your sources clearly.
            """,
            agent=agents["researcher"],
            context=[planning_task],
            expected_output="A comprehensive research report with factual information and citations",
            output_file="market_research.md"
        )
        
        fact_checking_task = Task(
            description="""
            Review the market research findings and verify the accuracy of key claims and data points.
            For any suspicious or unsupported claims:
            1. Identify the claim and why it might be questionable
            2. Conduct additional research to verify or correct it
            3. Provide correct information with proper citation
            
            Focus on the most important facts that would significantly impact the evaluation
            of the startup idea if they were inaccurate.
            """,
            agent=agents["fact_checker"],
            context=[research_task],
            expected_output="A verification report that confirms accurate information and corrects any inaccuracies",
            output_file="fact_check.md"
        )
        
        summary_task = Task(
            description="""
            Create a concise summary of the verified research findings.
            Structure your summary as follows:
            1. Key Market Statistics (3-5 bullet points)
            2. Major Trends (3-5 bullet points)
            3. Competitive Landscape (3-5 bullet points)
            4. Customer Needs Analysis (3-5 bullet points)
            
            Each bullet point should be clear, specific, and data-driven.
            Highlight the most important insights that would impact business decisions.
            """,
            agent=agents["summarizer"],
            context=[research_task, fact_checking_task],
            expected_output="A bulleted summary of key research findings",
            output_file="research_summary.md"
        )
        
        analysis_task = Task(
            description=f"""
            Analyze the business potential of the startup idea: '{startup_idea}'
            based on the research findings. Your analysis should include:
            
            1. SWOT Analysis (Strengths, Weaknesses, Opportunities, Threats)
            2. Competitive Advantage Assessment
            3. Market Fit Evaluation
            4. Business Model Viability
            5. Key Risk Factors
            
            For each component, provide specific recommendations or considerations
            that would improve the startup's chances of success.
            """,
            agent=agents["analyst"],
            context=[summary_task, research_task],
            expected_output="A detailed business analysis with specific recommendations",
            output_file="business_analysis.md"
        )
        
        strategy_task = Task(
            description="""
            Develop a go-to-market strategy recommendation based on the research and analysis.
            Your strategy should address:
            
            1. Target Customer Segmentation
            2. Value Proposition and Positioning
            3. Pricing Model Recommendations
            4. Distribution and Marketing Channels
            5. Key Partnerships and Ecosystem
            6. Initial Launch Approach
            
            Provide specific, actionable recommendations with justification based on the research.
            """,
            agent=agents["strategist"],
            context=[research_task, analysis_task],
            expected_output="A comprehensive go-to-market strategy with actionable recommendations",
            output_file="gtm_strategy.md"
        )
        
        report_writing_task = Task(
            description=f"""
            Create a comprehensive evaluation report for the startup idea: '{startup_idea}'
            based on all the research, analysis and strategy recommendations.
            
            Structure your report with the following sections:
            1. Executive Summary
            2. Market Overview and Opportunity
            3. Competitive Landscape
            4. Business Model Assessment
            5. SWOT Analysis
            6. Go-to-Market Strategy
            7. Risk Assessment
            8. Final Recommendations
            
            The report should be professional, evidence-based, and provide clear guidance
            on the viability of the startup idea and next steps.
            """,
            agent=agents["writer"],
            context=[summary_task, analysis_task, strategy_task],
            expected_output="A comprehensive startup evaluation report",
            output_file="startup_evaluation_report.md"
        )
        
        report_critique_task = Task(
            description="""
            Review the startup evaluation report and provide critical feedback to improve it.
            Focus on:
            
            1. Logical consistency and flow of arguments
            2. Evidence and support for key claims
            3. Clarity and precision of language
            4. Completeness of analysis
            5. Actionability of recommendations
            
            Provide specific suggestions for improvements, including areas that need
            more elaboration, clearer explanation, or stronger evidence.
            """,
            agent=agents["critic"],
            context=[report_writing_task],
            expected_output="A critique with specific improvement suggestions",
            output_file="report_critique.md"
        )
        
        final_report_task = Task(
            description="""
            Create the final version of the startup evaluation report by incorporating
            the critique and improvement suggestions.
            
            Ensure the final report is:
            1. Well-structured with clear section headings
            2. Professional in tone and presentation
            3. Evidence-based with specific data points
            4. Balanced in its assessment
            5. Actionable with clear recommendations
            
            This will be the final deliverable sent to the client.
            """,
            agent=agents["writer"],
            context=[report_writing_task, report_critique_task],
            expected_output="A polished final startup evaluation report",
            output_file="final_startup_evaluation.md"
        )
        
        return [
            planning_task,
            research_task,
            fact_checking_task,
            summary_task,
            analysis_task,
            strategy_task,
            report_writing_task,
            report_critique_task,
            final_report_task
        ]
    
    def evaluate_startup(self, startup_idea: str) -> Dict[str, Any]:
        """
        Run the full startup evaluation process.
        
        Args:
            startup_idea: The startup idea to evaluate
            
        Returns:
            A dictionary containing the results and file paths
        """
        # Generate a unique ID for this research
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        idea_slug = startup_idea.lower().replace(" ", "_")[:30]
        research_id = f"{idea_slug}_{timestamp}"
        
        logger.info(f"Starting evaluation of: {startup_idea} (ID: {research_id})")
        
        # Create the agents and tasks
        agents = self.create_agents()
        tasks = self.create_tasks(agents, startup_idea)
        
        # Create and run the crew
        crew = Crew(
            agents=list(agents.values()),
            tasks=tasks,
            verbose=True,
            process=self.process
        )
        
        try:
            # Run the evaluation
            result = crew.kickoff()
            
            # Create directory for this specific research
            research_dir = self.output_dir / research_id
            os.makedirs(research_dir, exist_ok=True)
            
            # Save all task outputs
            files = {}
            for task in tasks:
                output_path = research_dir / task.output_file
                with open(output_path, 'w', encoding='utf-8') as f:
                    if isinstance(task.output, TaskOutput):
                        f.write(str(task.output))
                    elif task.output:
                        f.write(str(task.output))
                files[task.output_file] = str(output_path)
            
            # Generate additional output formats for the final report
            final_report_path = research_dir / "final_startup_evaluation.md"
            
            # Create PDF (simulated as HTML for now)
            html_path = research_dir / "final_startup_evaluation.html"
            if os.path.exists(final_report_path):
                with open(final_report_path, 'r', encoding='utf-8') as f:
                    markdown_content = f.read()
                html_content = markdown2.markdown(markdown_content)
                with open(html_path, 'w', encoding='utf-8') as f:
                    f.write(html_content)
                files["html"] = str(html_path)
            
            # Create DOCX
            docx_path = research_dir / "final_startup_evaluation.docx"
            if os.path.exists(final_report_path):
                doc = Document()
                doc.add_heading(f"Startup Evaluation: {startup_idea}", 0)
                
                with open(final_report_path, 'r', encoding='utf-8') as f:
                    markdown_content = f.read()
                
                # Simple markdown parsing (headers and paragraphs only)
                for line in markdown_content.split('\n'):
                    if line.startswith('# '):
                        doc.add_heading(line[2:], 1)
                    elif line.startswith('## '):
                        doc.add_heading(line[3:], 2)
                    elif line.startswith('### '):
                        doc.add_heading(line[4:], 3)
                    elif line.strip():
                        doc.add_paragraph(line)
                
                doc.save(docx_path)
                files["docx"] = str(docx_path)
            
            # Return results
            return {
                "research_id": research_id,
                "startup_idea": startup_idea,
                "status": "completed",
                "files": files,
                "summary": result
            }
            
        except Exception as e:
            logger.error(f"Error during startup evaluation: {str(e)}", exc_info=True)
            return {
                "research_id": research_id,
                "startup_idea": startup_idea,
                "status": "error",
                "error": str(e)
            }
    
    def get_research_by_id(self, research_id: str) -> Dict[str, Any]:
        """Retrieve a previously completed research by its ID"""
        research_dir = self.output_dir / research_id
        
        if not os.path.exists(research_dir):
            return {"status": "not_found", "error": "Research ID not found"}
        
        files = {}
        for file_path in research_dir.glob("*"):
            files[file_path.name] = str(file_path)
        
        return {
            "research_id": research_id,
            "status": "completed",
            "files": files
        }
    
    def list_researches(self) -> List[Dict[str, str]]:
        """List all completed researches"""
        researches = []
        
        for research_dir in self.output_dir.glob("*"):
            if research_dir.is_dir():
                # Try to extract the startup idea from the directory name
                parts = research_dir.name.split("_")
                timestamp_parts = [p for p in parts if p.isdigit() and len(p) >= 8]
                if timestamp_parts:
                    # Remove timestamp parts from name
                    for tp in timestamp_parts:
                        parts.remove(tp)
                    startup_idea = " ".join(parts).replace("_", " ").title()
                else:
                    startup_idea = research_dir.name.replace("_", " ").title()
                
                researches.append({
                    "research_id": research_dir.name,
                    "startup_idea": startup_idea,
                    "created_at": datetime.datetime.fromtimestamp(research_dir.stat().st_ctime).isoformat()
                })
        
        # Sort by creation time (newest first)
        researches.sort(key=lambda x: x["created_at"], reverse=True)
        return researches


# For direct testing
if __name__ == "__main__":
    # Set your API key
    # os.environ["OPENAI_API_KEY"] = "your-api-key"
    
    crew = StartupResearchCrew()
    result = crew.evaluate_startup("AI tool for personalized mental health")
    print(json.dumps(result, indent=2)) 